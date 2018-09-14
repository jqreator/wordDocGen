#!/usr/bin/env python
#Word Document Generator v1.0.0
#Copyright (C)2018 Jacob Parks - jqreator at gmail dot com
#
#This program is free software: you can redistribute it and/or modify
#it under the terms of the GNU General Public License as published by
#the Free Software Foundation, either version 3 of the License, or
#at your option) any later version.
#
#This program is distributed in the hope that it will be useful,
#but WITHOUT ANY WARRANTY; without even the implied warranty of
#MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
#GNU General Public License for more details.
#
#You should have received a copy of the GNU General Public License
#along with this program.  If not, see <http://www.gnu.org/licenses/>.
#
#Release notes:
#v1.0.0 - Initial release

#Smallest file is 36KB

try:
    import docx
except ImportError:
    print('\nThe python-docx module does not appear to be installed on this system.')
    print('\nTo install with pip, use pip install python-docx')

import docx
import random
import string
import os

docName = 'doc'
numFiles = int(input('How many documents would you like to create?'))
MB = 1350000 #1MB is around 1350000 characters
charLen = MB * 1
doc = docx.Document()
jibberish = random.choices(string.ascii_uppercase + string.ascii_lowercase + string.digits, k=charLen)

if not os.path.exists('testdocs'):
    print('Creating "testdocs" directory to store files')
    os.mkdir('testdocs')

print('Creating ' + str(numFiles) + ' files...')

for x in range(0, numFiles):
    doc = docx.Document()
    content = []
    content = jibberish
    content.append(str(x))
    doc.add_paragraph(content)
    doc.save('testdocs\\' + docName + str(x) + '.docx')
    x -= 1

print()
print('Done.')
print()
input('Press and key to quit: ')

